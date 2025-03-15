from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading('Inference Module Architecture for AD-CNN and FD-CNN', level=1)
doc.add_paragraph(
    "The following tables summarize the layer-by-layer architecture for different target sizes used in the inference module. "
    "Each table lists the layers in order along with their key parameters."
)

def add_table(doc, target_size, rows_data):
    doc.add_heading(f'Target Size: {target_size}', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer Type'
    hdr_cells[1].text = 'Parameters / Details'
    for row in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
    doc.add_paragraph("")  # Add a blank line

# Data for each target size
data_8 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 4, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 2, Stride = 2, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 64, Out Features = 1")
]

data_16 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 4, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 2, Stride = 2, Padding = 0"),
    ("Conv1d", "In: 32 channels, Out: 64 filters, Kernel Size = 2, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 2, Stride = 2, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 128, Out Features = 1")
]

data_32 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 8, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 4, Stride = 4, Padding = 0"),
    ("Conv1d", "In: 32 channels, Out: 64 filters, Kernel Size = 2, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 2, Stride = 2, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 128, Out Features = 1")
]

data_64 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 8, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 4, Stride = 4, Padding = 0"),
    ("Conv1d", "In: 32 channels, Out: 64 filters, Kernel Size = 4, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 4, Stride = 4, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 128, Out Features = 1")
]

data_128 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 16, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 8, Stride = 4, Padding = 0"),
    ("Conv1d", "In: 32 channels, Out: 64 filters, Kernel Size = 8, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 4, Stride = 4, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 320, Out Features = 64"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("Linear", "In Features = 64, Out Features = 1")
]

data_256 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 16, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 8, Stride = 6, Padding = 0"),
    ("Conv1d", "In: 32 channels, Out: 64 filters, Kernel Size = 8, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 4, Stride = 4, Padding = 0"),
    ("Conv1d", "In: 64 channels, Out: 128 filters, Kernel Size = 4, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "128 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 2, Stride = 2, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 256, Out Features = 1")
]
data_512 = [
    ("Conv1d", "In: 1 channel, Out: 32 filters, Kernel Size = 16, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 8, Stride = 8, Padding = 0"),
    ("Conv1d", "In: 32 channels, Out: 64 filters, Kernel Size = 8, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 8, Stride = 8, Padding = 0"),
    ("Conv1d", "In: 64 channels, Out: 128 filters, Kernel Size = 4, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "128 channels"),
    ("Activation", "LeakyReLU"),
    ("MaxPool1d", "Kernel Size = 2, Stride = 2, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "In Features = 128, Out Features = 1")
]

add_table(doc, 8, data_8)
add_table(doc, 16, data_16)
add_table(doc, 32, data_32)
add_table(doc, 64, data_64)
add_table(doc, 128, data_128)
add_table(doc, 256, data_256)
add_table(doc, 512, data_512)

doc.save("InferenceModuleArchitecture.docx")
