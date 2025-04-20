import numpy as np
from docx import Document
import final_selected_bands as fsd


def get_mid_bands(size):
    x = np.linspace(0, 4199, size+2)
    x = x[1:-1]
    x = np.round(x)
    x = x.astype(int)
    x = [fsd.band_to_wl(i) for i in x]
    return x




document = Document()
table = document.add_table(rows=1, cols=2)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Lower dimensional size'
hdr_cells[1].text = 'Selected bands'

for key, value in fsd.selected_wl_as_list.items():
    ad_wl =  list(dict.fromkeys(value))
    row_cells = table.add_row().cells
    row_cells[0].text = str(key)
    row_cells[1].text = ', '.join(map(str, ad_wl))
    document.save('band_list_ad.docx')

document = Document()
table = document.add_table(rows=1, cols=2)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Lower dimensional size'
hdr_cells[1].text = 'Selected bands'
for key, value in fsd.bsdr_wl_as_list.items():
    ad_wl =  list(dict.fromkeys(value))
    row_cells = table.add_row().cells
    row_cells[0].text = str(key)
    row_cells[1].text = ', '.join(map(str, ad_wl))
    document.save('band_list_bd.docx')



document = Document()
table = document.add_table(rows=1, cols=2)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Lower dimensional size'
hdr_cells[1].text = 'Selected bands'
for key, value in fsd.bsdr_wl_as_list.items():
    ad_wl = get_mid_bands(key)
    row_cells = table.add_row().cells
    row_cells[0].text = str(key)
    row_cells[1].text = ', '.join(map(str, ad_wl))
    document.save('band_list_fd.docx')