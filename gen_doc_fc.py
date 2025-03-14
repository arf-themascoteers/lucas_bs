from docx import Document
from docx.shared import Pt

# Create a new document
doc = Document()
doc.add_heading('Fully-Connected Module Architecture', level=1)
doc.add_paragraph(
    "The following tables summarize the architecture of the fully-connected (FC) module used in the network. "
    "For target sizes in the range (100, 512], two hidden layers are used with sizes target_size//2 and target_size//4 "
    "respectively, each followed by BatchNorm1d and LeakyReLU. For target sizes <= 100, a single hidden layer of size "
    "max(10, target_size//2) is used."
)

def add_table(doc, target_size, rows):
    # Add a subheading for the target size
    doc.add_heading(f"Target Size: {target_size}", level=2)
    # Create a table with 2 columns: one for the layer and one for its details.
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Layer"
    hdr_cells[1].text = "Parameters / Details"
    for layer, details in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = details
    doc.add_paragraph("")  # Blank line for spacing

def fc_architecture_large(target_size):
    hidden1 = target_size // 2
    hidden2 = target_size // 4
    return [
        ("Linear", f"In Features: {target_size}, Out Features: {hidden1}"),
        ("BatchNorm1d", f"{hidden1} channels"),
        ("Activation", "LeakyReLU"),
        ("Linear", f"In Features: {hidden1}, Out Features: {hidden2}"),
        ("BatchNorm1d", f"{hidden2} channels"),
        ("Activation", "LeakyReLU"),
        ("Linear", f"In Features: {hidden2}, Out Features: 1")
    ]

def fc_architecture_small(target_size):
    hidden = max(10, target_size // 2)
    return [
        ("Linear", f"In Features: {target_size}, Out Features: {hidden}"),
        ("BatchNorm1d", f"{hidden} channels"),
        ("Activation", "LeakyReLU"),
        ("Linear", f"In Features: {hidden}, Out Features: 1")
    ]

# Define target sizes and generate table data accordingly.
# For target sizes > 100 and <= 512, we use the 'large' architecture.
# For target sizes <= 100, we use the 'small' architecture.
target_sizes = [8, 16, 32, 64, 128, 256, 512]

for ts in target_sizes:
    if 100 < ts <= 512:
        rows = fc_architecture_large(ts)
    else:
        rows = fc_architecture_small(ts)
    add_table(doc, ts, rows)

# Save the document to a file.
doc.save("FCModuleArchitecture.docx")
