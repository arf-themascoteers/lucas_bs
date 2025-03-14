from docx import Document
from docx.shared import Pt

# Create a new document
doc = Document()
doc.add_heading('ANN Architecture', level=1)
doc.add_paragraph(
    "The following table documents the structure of the 1D-CNN used in our study. "
    "The architecture is defined as follows:"
)

# Create the table with 2 columns: Layer and Parameters / Details
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Layer"
hdr_cells[1].text = "Parameters / Details"

# Define the layers and their details
layers = [
    ("Conv1d", "Input: 1 channel, Output: 32 filters, Kernel Size = 16, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "32 channels"),
    ("LeakyReLU", ""),
    ("MaxPool1d", "Kernel Size = 16, Stride = 16, Padding = 0"),
    ("Conv1d", "Input: 32 channels, Output: 64 filters, Kernel Size = 8, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "64 channels"),
    ("LeakyReLU", ""),
    ("MaxPool1d", "Kernel Size = 16, Stride = 16, Padding = 0"),
    ("Conv1d", "Input: 64 channels, Output: 128 filters, Kernel Size = 4, Stride = 1, Padding = 0"),
    ("BatchNorm1d", "128 channels"),
    ("LeakyReLU", ""),
    ("MaxPool1d", "Kernel Size = 8, Stride = 8, Padding = 0"),
    ("Flatten", "Starting from dimension 1"),
    ("Linear", "Input Features: 128, Output Features: 1")
]

# Populate the table with layer details
for layer, details in layers:
    row_cells = table.add_row().cells
    row_cells[0].text = layer
    row_cells[1].text = details

# Save the document to a DOCX file
doc.save("ANNArchitecture.docx")
